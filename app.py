import os
import re
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from lxml.etree import QName
from openpyxl import load_workbook
from openpyxl.styles import Font

UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "output"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

WORD_NAMESPACE = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

def secure_thai_filename(filename):
    filename = filename.strip().replace(" ", "_")
    filename = re.sub(r'[^\wก-๙_.()-]', '', filename)  
    return filename

def set_word_font(run, font_name="Times New Roman", font_size=12):
    run.font.name = font_name
    run.font.size = Pt(font_size)

    r = run._element
    rPr = r.find(f".//{{{WORD_NAMESPACE}}}rPr")

    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.append(rPr)

    rFonts = rPr.find(f".//{{{WORD_NAMESPACE}}}rFonts")
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)

    rFonts.set(QName(WORD_NAMESPACE, "ascii"), font_name)
    rFonts.set(QName(WORD_NAMESPACE, "hAnsi"), font_name)
    rFonts.set(QName(WORD_NAMESPACE, "cs"), font_name)

def change_word_font(doc_path, output_path):
    doc = Document(doc_path)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            set_word_font(run)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        set_word_font(run)

    doc.save(output_path)
    print(f"Processed and saved Word file: {output_path}")

def change_excel_font(excel_path, output_path):
    try:
        wb = load_workbook(excel_path, data_only=True)
        print(f"Opened Excel file: {excel_path}")

        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value:
                        cell.font = Font(name="Times New Roman", size=12)

        wb.save(output_path)
        print(f"Processed and saved Excel file: {output_path}")

    except Exception as e:
        print(f"Error processing {excel_path}: {e}")

def process_files():
    files = os.listdir(UPLOAD_FOLDER)
    if not files:
        print("No files found in the uploads directory.")
        return

    for filename in files:
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file_ext = os.path.splitext(filename)[1].lower()
        
        if file_ext == ".docx":
            new_filename = f"{os.path.splitext(filename)[0]}-word.docx"
            output_path = os.path.join(OUTPUT_FOLDER, new_filename)
            change_word_font(file_path, output_path)

        elif file_ext == ".xlsx":
            new_filename = f"{os.path.splitext(filename)[0]}-excel.xlsx"
            output_path = os.path.join(OUTPUT_FOLDER, new_filename)
            change_excel_font(file_path, output_path)

if __name__ == "__main__":
    process_files()
